#!/usr/bin/env python3
"""
Advanced Document Evaluation MCP Server
Superior implementation with enhanced features
"""

import json
import logging
from pathlib import Path
from typing import Dict, List, Any, Optional, Union
from dataclasses import dataclass
from enum import Enum
import pandas as pd
from pydantic import BaseModel, Field

from mcp.server.fastmcp import FastMCP

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Initialize FastMCP server
server = FastMCP("DocEval")

class SupportedFormat(Enum):
    """Supported document formats"""
    PDF = "pdf"
    DOC = "doc"
    DOCX = "docx"
    TXT = "txt"
    RTF = "rtf"
    HTML = "html"
    MD = "md"

class CriteriaFormat(Enum):
    """Supported criteria formats"""
    JSON = "json"
    CSV = "csv"
    XLSX = "xlsx"
    YAML = "yaml"

@dataclass
class ParsedDocument:
    """Structured document data"""
    content: str
    metadata: Dict[str, Any]
    word_count: int
    char_count: int
    format: SupportedFormat

class Criterion(BaseModel):
    """Individual evaluation criterion"""
    name: str = Field(..., description="Criterion name")
    description: str = Field(..., description="Detailed description") 
    weight: float = Field(default=1.0, ge=0.1, le=10.0, description="Weight 0.1-10.0")
    max_score: int = Field(default=10, ge=1, le=100, description="Maximum possible score")
    category: Optional[str] = Field(None, description="Category grouping")

class EvaluationCriteria(BaseModel):
    """Complete criteria set"""
    criteria: List[Criterion]
    metadata: Dict[str, Any] = Field(default_factory=dict)
    total_weight: Optional[float] = None

class ScoreResult(BaseModel):
    """Individual score result"""
    criterion: str
    score: float
    max_score: int
    percentage: float
    justification: str
    confidence: float = Field(ge=0.0, le=1.0)
    category: Optional[str] = None

class EvaluationResult(BaseModel):
    """Complete evaluation result"""
    scores: List[ScoreResult]
    overall_score: float
    weighted_score: float
    max_possible: float
    percentage: float
    summary: str
    recommendations: List[str] = Field(default_factory=list)
    metadata: Dict[str, Any] = Field(default_factory=dict)

def parse_document_content(file_path: str) -> ParsedDocument:
    """Enhanced document parsing with metadata extraction"""
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {file_path}")
    
    file_format = SupportedFormat(path.suffix.lower().lstrip('.'))
    metadata = {
        "file_size": path.stat().st_size,
        "file_name": path.name,
        "extension": path.suffix,
        "modified_time": path.stat().st_mtime
    }
    
    try:
        if file_format == SupportedFormat.PDF:
            content = _parse_pdf_advanced(path)
        elif file_format in [SupportedFormat.DOC, SupportedFormat.DOCX]:
            content = _parse_word_document(path)
        elif file_format == SupportedFormat.TXT:
            content = path.read_text(encoding='utf-8')
        elif file_format == SupportedFormat.MD:
            content = path.read_text(encoding='utf-8')
        else:
            # Fallback to textract for other formats
            import textract
            content = textract.process(str(path)).decode('utf-8')
            
    except Exception as e:
        logger.error(f"Error parsing {file_path}: {e}")
        raise ValueError(f"Could not parse document: {e}")
    
    return ParsedDocument(
        content=content.strip(),
        metadata=metadata,
        word_count=len(content.split()),
        char_count=len(content),
        format=file_format
    )

def _parse_pdf_advanced(path: Path) -> str:
    """Advanced PDF parsing with better text extraction"""
    import pypdf
    
    content = ""
    with open(path, 'rb') as file:
        reader = pypdf.PdfReader(file)
        
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    content += f"\n--- Page {page_num + 1} ---\n{page_text}"
            except Exception as e:
                logger.warning(f"Could not extract text from page {page_num + 1}: {e}")
                
    return content.strip()

def _parse_word_document(path: Path) -> str:
    """Parse Word documents with python-docx"""
    from docx import Document
    
    doc = Document(path)
    content = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            content.append(paragraph.text)
    
    return "\n".join(content)

@server.tool()
def parse_document(file_path: str) -> Dict[str, Any]:
    """
    Advanced document parsing with metadata extraction.
    
    Args:
        file_path: Path to document file
        
    Returns:
        Parsed document with content and metadata
    """
    try:
        parsed = parse_document_content(file_path)
        return {
            "content": parsed.content,
            "metadata": parsed.metadata,
            "word_count": parsed.word_count,
            "char_count": parsed.char_count,
            "format": parsed.format.value,
            "success": True
        }
    except Exception as e:
        logger.error(f"Document parsing failed: {e}")
        return {"error": str(e), "success": False}

@server.tool()
def parse_criteria(criteria_path: str) -> Dict[str, Any]:
    """
    Advanced criteria parsing supporting multiple formats.
    
    Args:
        criteria_path: Path to criteria file (JSON, CSV, XLSX, YAML)
        
    Returns:
        Structured criteria with validation
    """
    try:
        path = Path(criteria_path)
        format_type = CriteriaFormat(path.suffix.lower().lstrip('.'))
        
        if format_type == CriteriaFormat.JSON:
            criteria_data = _parse_json_criteria(path)
        elif format_type == CriteriaFormat.CSV:
            criteria_data = _parse_csv_criteria(path)
        elif format_type == CriteriaFormat.XLSX:
            criteria_data = _parse_xlsx_criteria(path)
        else:
            raise ValueError(f"Unsupported criteria format: {format_type}")
        
        # Validate with Pydantic
        criteria = EvaluationCriteria(**criteria_data)
        criteria.total_weight = sum(c.weight for c in criteria.criteria)
        
        return {
            "criteria": [c.model_dump() for c in criteria.criteria],
            "total_weight": criteria.total_weight,
            "count": len(criteria.criteria),
            "metadata": criteria.metadata,
            "success": True
        }
        
    except Exception as e:
        logger.error(f"Criteria parsing failed: {e}")
        return {"error": str(e), "success": False}

def _parse_json_criteria(path: Path) -> Dict[str, Any]:
    """Parse JSON criteria with validation"""
    with open(path, 'r') as f:
        data = json.load(f)
    
    if "criteria" not in data:
        # Auto-convert simple format
        data = {"criteria": data}
    
    return data

def _parse_csv_criteria(path: Path) -> Dict[str, Any]:
    """Parse CSV criteria with pandas"""
    df = pd.read_csv(path)
    
    # Expected columns: name, description, weight, max_score, category
    required_cols = ['name', 'description']
    if not all(col in df.columns for col in required_cols):
        # Try alternative column names
        col_mapping = {
            'criterion': 'name',
            'criteria': 'name', 
            'desc': 'description',
            'weight': 'weight',
            'max': 'max_score',
            'category': 'category'
        }
        df = df.rename(columns=col_mapping)
    
    criteria = []
    for _, row in df.iterrows():
        criterion = {
            "name": str(row.get('name', '')),
            "description": str(row.get('description', '')),
            "weight": float(row.get('weight', 1.0)),
            "max_score": int(row.get('max_score', 10)),
            "category": str(row.get('category', '')) if pd.notna(row.get('category')) else None
        }
        criteria.append(criterion)
    
    return {"criteria": criteria}

def _parse_xlsx_criteria(path: Path) -> Dict[str, Any]:
    """Parse Excel criteria"""
    df = pd.read_excel(path)
    return _parse_csv_criteria(path)  # Same logic as CSV

@server.tool()
def evaluate_document_advanced(
    document_content: str,
    criteria_data: Dict[str, Any],
    model_name: str = "llama3.2",
    evaluation_style: str = "comprehensive"
) -> Dict[str, Any]:
    """
    Advanced AI-powered document evaluation with multiple analysis modes.
    
    Args:
        document_content: Document text content
        criteria_data: Parsed criteria structure
        model_name: Ollama model to use
        evaluation_style: "comprehensive", "quick", "detailed"
        
    Returns:
        Structured evaluation results with scores and analysis
    """
    try:
        import ollama
        
        criteria = criteria_data.get("criteria", [])
        if not criteria:
            return {"error": "No criteria provided", "success": False}
        
        # Build sophisticated evaluation prompt
        prompt = _build_evaluation_prompt(document_content, criteria, evaluation_style)
        
        response = ollama.chat(
            model=model_name,
            messages=[{"role": "user", "content": prompt}]
        )
        
        # Parse and validate response
        result = _parse_evaluation_response(response["message"]["content"], criteria)
        
        return {
            "evaluation": result.model_dump(),
            "model_used": model_name,
            "style": evaluation_style,
            "success": True
        }
        
    except Exception as e:
        logger.error(f"Evaluation failed: {e}")
        return {"error": str(e), "success": False}

def _build_evaluation_prompt(content: str, criteria: List[Dict], style: str) -> str:
    """Build sophisticated evaluation prompt"""
    
    style_instructions = {
        "comprehensive": "Provide thorough analysis with detailed justifications",
        "quick": "Focus on key points with concise explanations", 
        "detailed": "Include specific examples and comprehensive recommendations"
    }
    
    criteria_text = ""
    for i, criterion in enumerate(criteria, 1):
        criteria_text += f"""
{i}. **{criterion['name']}** (Weight: {criterion.get('weight', 1.0)}, Max: {criterion.get('max_score', 10)})
   Description: {criterion['description']}
   Category: {criterion.get('category', 'General')}
"""

    content_preview = content[:1500] + "..." if len(content) > 1500 else content
    
    prompt = f"""
You are an expert document evaluator. Analyze the following document against the specified criteria.

EVALUATION STYLE: {style} - {style_instructions.get(style, '')}

DOCUMENT TO EVALUATE:
{content_preview}

EVALUATION CRITERIA:
{criteria_text}

INSTRUCTIONS:
1. Score each criterion on its specified scale (default 1-{criteria[0].get('max_score', 10)})
2. Provide specific justification for each score
3. Include confidence level (0.0-1.0) for each assessment
4. Give actionable recommendations for improvement
5. Consider the document's purpose and audience

REQUIRED JSON RESPONSE FORMAT:
{{
    "scores": [
        {{
            "criterion": "criterion_name",
            "score": numeric_score,
            "max_score": max_possible,
            "percentage": percentage_value,
            "justification": "detailed_explanation",
            "confidence": confidence_level,
            "category": "category_name"
        }}
    ],
    "overall_score": weighted_average,
    "weighted_score": calculated_weighted_score,
    "max_possible": total_max_score,
    "percentage": overall_percentage,
    "summary": "comprehensive_summary",
    "recommendations": ["specific_improvement_1", "specific_improvement_2"]
}}

Ensure all scores are numeric and within the specified ranges.
"""
    
    return prompt

def _parse_evaluation_response(response: str, criteria: List[Dict]) -> EvaluationResult:
    """Parse and validate evaluation response"""
    try:
        # Extract JSON from response
        start = response.find('{')
        end = response.rfind('}') + 1
        if start == -1 or end == 0:
            raise ValueError("No JSON found in response")
            
        json_str = response[start:end]
        data = json.loads(json_str)
        
        # Validate and structure result
        return EvaluationResult(**data)
        
    except (json.JSONDecodeError, ValueError) as e:
        logger.warning(f"Could not parse structured response: {e}")
        
        # Fallback: create basic result
        scores = []
        for criterion in criteria:
            scores.append(ScoreResult(
                criterion=criterion['name'],
                score=5.0,  # Default mid-range score
                max_score=criterion.get('max_score', 10),
                percentage=50.0,
                justification=f"Could not parse detailed evaluation for {criterion['name']}",
                confidence=0.3,
                category=criterion.get('category')
            ))
        
        return EvaluationResult(
            scores=scores,
            overall_score=5.0,
            weighted_score=5.0,
            max_possible=10.0,
            percentage=50.0,
            summary="Evaluation completed but response parsing failed",
            recommendations=["Retry evaluation with different model"]
        )

@server.tool()
def calculate_advanced_metrics(evaluation_result: Dict[str, Any]) -> Dict[str, Any]:
    """
    Calculate advanced metrics and insights from evaluation results.
    
    Args:
        evaluation_result: Results from evaluate_document_advanced
        
    Returns:
        Advanced metrics and analysis
    """
    try:
        eval_data = evaluation_result.get("evaluation", {})
        scores = eval_data.get("scores", [])
        
        if not scores:
            return {"error": "No scores to analyze", "success": False}
        
        # Calculate various metrics
        score_values = [s["score"] for s in scores]
        max_scores = [s["max_score"] for s in scores]
        weights = [s.get("weight", 1.0) for s in scores]
        
        metrics = {
            "statistical_analysis": {
                "mean": sum(score_values) / len(score_values),
                "median": sorted(score_values)[len(score_values) // 2],
                "std_dev": _calculate_std_dev(score_values),
                "min_score": min(score_values),
                "max_score": max(score_values),
                "range": max(score_values) - min(score_values)
            },
            "performance_bands": _categorize_scores(scores),
            "improvement_priorities": _identify_priorities(scores),
            "strength_areas": _identify_strengths(scores),
            "risk_assessment": _assess_risks(scores),
            "consistency_score": _calculate_consistency(score_values),
            "success": True
        }
        
        return metrics
        
    except Exception as e:
        logger.error(f"Metrics calculation failed: {e}")
        return {"error": str(e), "success": False}

def _calculate_std_dev(values: List[float]) -> float:
    """Calculate standard deviation"""
    if len(values) < 2:
        return 0.0
    mean = sum(values) / len(values)
    variance = sum((x - mean) ** 2 for x in values) / (len(values) - 1)
    return variance ** 0.5

def _categorize_scores(scores: List[Dict]) -> Dict[str, List[str]]:
    """Categorize scores into performance bands"""
    bands = {"excellent": [], "good": [], "fair": [], "poor": []}
    
    for score in scores:
        percentage = score.get("percentage", 0)
        criterion = score.get("criterion", "Unknown")
        
        if percentage >= 90:
            bands["excellent"].append(criterion)
        elif percentage >= 75:
            bands["good"].append(criterion)
        elif percentage >= 60:
            bands["fair"].append(criterion)
        else:
            bands["poor"].append(criterion)
    
    return bands

def _identify_priorities(scores: List[Dict]) -> List[Dict[str, Any]]:
    """Identify improvement priorities"""
    priorities = []
    
    for score in scores:
        if score.get("percentage", 100) < 70:  # Below acceptable threshold
            priorities.append({
                "criterion": score.get("criterion"),
                "current_score": score.get("score"),
                "potential_gain": score.get("max_score", 10) - score.get("score", 0),
                "priority_level": "high" if score.get("percentage", 100) < 50 else "medium"
            })
    
    return sorted(priorities, key=lambda x: x["potential_gain"], reverse=True)

def _identify_strengths(scores: List[Dict]) -> List[str]:
    """Identify strength areas"""
    return [
        score.get("criterion", "Unknown") 
        for score in scores 
        if score.get("percentage", 0) >= 85
    ]

def _assess_risks(scores: List[Dict]) -> Dict[str, Any]:
    """Assess risks based on low scores"""
    critical_low = [s for s in scores if s.get("percentage", 100) < 40]
    concerning_low = [s for s in scores if 40 <= s.get("percentage", 100) < 60]
    
    return {
        "critical_issues": len(critical_low),
        "concerning_areas": len(concerning_low),
        "risk_level": "high" if critical_low else "medium" if concerning_low else "low",
        "critical_criteria": [s.get("criterion") for s in critical_low]
    }

def _calculate_consistency(scores: List[float]) -> float:
    """Calculate consistency score (inverse of coefficient of variation)"""
    if not scores or len(scores) < 2:
        return 1.0
    
    mean_score = sum(scores) / len(scores)
    if mean_score == 0:
        return 0.0
        
    std_dev = _calculate_std_dev(scores)
    cv = std_dev / mean_score  # Coefficient of variation
    
    # Convert to consistency score (0-1, higher is more consistent)
    return max(0, 1 - cv)

@server.tool()
def full_document_evaluation(
    document_path: str,
    criteria_path: str,
    model_name: str = "llama3.2",
    evaluation_style: str = "comprehensive"
) -> Dict[str, Any]:
    """
    Complete end-to-end document evaluation workflow.
    
    Args:
        document_path: Path to document file
        criteria_path: Path to criteria file
        model_name: Ollama model to use
        evaluation_style: Evaluation approach
        
    Returns:
        Complete evaluation with all metrics and analysis
    """
    try:
        # Step 1: Parse document
        doc_result = parse_document(document_path)
        if not doc_result.get("success"):
            return doc_result
        
        # Step 2: Parse criteria
        criteria_result = parse_criteria(criteria_path)
        if not criteria_result.get("success"):
            return criteria_result
        
        # Step 3: Evaluate document
        eval_result = evaluate_document_advanced(
            doc_result["content"],
            criteria_result,
            model_name,
            evaluation_style
        )
        if not eval_result.get("success"):
            return eval_result
        
        # Step 4: Calculate advanced metrics
        metrics_result = calculate_advanced_metrics(eval_result)
        
        # Combine all results
        return {
            "document_info": {
                "path": document_path,
                "word_count": doc_result["word_count"],
                "char_count": doc_result["char_count"],
                "format": doc_result["format"]
            },
            "criteria_info": {
                "path": criteria_path,
                "count": criteria_result["count"],
                "total_weight": criteria_result["total_weight"]
            },
            "evaluation": eval_result["evaluation"],
            "advanced_metrics": metrics_result,
            "model_used": model_name,
            "evaluation_style": evaluation_style,
            "success": True
        }
        
    except Exception as e:
        logger.error(f"Full evaluation failed: {e}")
        return {"error": str(e), "success": False}

def main():
    """Start the MCP server"""
    logger.info("Starting DocEval MCP Server...")
    server.run()

if __name__ == "__main__":
    main()